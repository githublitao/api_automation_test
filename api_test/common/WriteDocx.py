import json
import logging

import docx
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

logger = logging.getLogger(__name__)  # 这里使用 __name__ 动态搜索定义的 logger 配置，这里有一个层次关系的知识点。


class Write:
    def __init__(self):     
        self.doc = docx.Document()
        self.doc.styles['Normal'].font.name = u'宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def write_api(self, api_name, group_data=None, data=None):
        self.doc.add_paragraph(style="Title").add_run(api_name)
        index = 1
        if group_data:
            for item in group_data:
                self.doc.add_paragraph(style="Heading 1").add_run(str(index)+"、"+item["name"]).font.size = 300000
                _id = 1
                for items in item['First']:
                    self.doc.add_paragraph(style="Heading 2").add_run(str(_id)+"."+items['name'],
                                                                      style="Default Paragraph Font")
                    _id = _id+1
                    self.doc.add_paragraph("简要描述：", style="Body Text").paragraph_format.space_before = Pt(14)
                    self.doc.add_paragraph(items['description'], style="List Paragraph")
                    self.doc.add_paragraph("请求URL：", style="Body Text")
                    if items['httpType'] == 'HTTP':
                        self.doc.add_paragraph("http://xxxx"+items['apiAddress'], style="List Paragraph")
                    else:
                        self.doc.add_paragraph("https://xxxx" + items['apiAddress'], style="List Paragraph")
                    self.doc.add_paragraph("请求方式：", style="Body Text")
                    self.doc.add_paragraph(items['requestType'], style="List Paragraph")
                    self.doc.add_paragraph("参数：", style="Body Text")
                    if items['requestParameterType'] == 'form-data':
                        table = self.doc.add_table(rows=len(items['requestParameter'])+1,
                                                   cols=4, style="Medium Shading 1 Accent 1")
                        hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
                        hdr_cells[0].text = '参数名'  # 第一行的第一列,给这行里面添加文字
                        hdr_cells[1].text = '参数类型'
                        hdr_cells[2].text = '必填?'
                        hdr_cells[3].text = '输入限制'
                        j = 0
                        for row in table.rows[1:]:
                            row.cells[0].text = items['requestParameter'][j]['name']
                            row.cells[1].text = items['requestParameter'][j]['_type']
                            if items['requestParameter'][j]['required']:
                                row.cells[2].text = "是"
                            else:
                                row.cells[2].text = "否"
                            try:
                                row.cells[3].text = items['requestParameter'][j]['restrict']
                            except TypeError:
                                pass
                            j = j+1
                    else:
                        try:
                            if len(items['requestParameterRaw']):
                                data = json.loads(items['requestParameterRaw'][0]["data"])
                                self.doc.add_paragraph(style="Normal").add_run('{')
                                write_json(self.doc, data, 0.3)
                                self.doc.add_paragraph(style="Normal").add_run('}')
                        except:
                            logging.exception("Error")
                    self.doc.add_paragraph("返回参数：", style="Body Text")
                    table = self.doc.add_table(rows=len(items['response']) + 1,
                                               cols=4, style="Medium Shading 1 Accent 1")
                    hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
                    hdr_cells[0].text = '参数名'  # 第一行的第一列,给这行里面添加文字
                    hdr_cells[1].text = '参数类型'
                    hdr_cells[2].text = '必含?'
                    hdr_cells[3].text = '说明'
                    j = 0
                    for row in table.rows[1:]:
                        row.cells[0].text = items['response'][j]['name']
                        row.cells[1].text = items['response'][j]['_type']
                        if items['response'][j]['required']:
                            row.cells[2].text = "是"
                        else:
                            row.cells[2].text = "否"
                        try:
                            row.cells[3].text = items['response'][j]['description']
                        except TypeError:
                            pass
                        j = j + 1
                    self.doc.add_paragraph()
                    self.doc.add_paragraph("返回示例：", style="Body Text")
                    try:
                        if len(items['data']):
                            data = eval(items['data'].replace("true", "True").replace("false", "False").replace("null", "None"))
                            self.doc.add_paragraph(style="Normal").add_run('{')
                            write_json(self.doc, data, 0.3)
                            self.doc.add_paragraph(style="Normal").add_run('}')
                    except:
                        logging.exception("Error")
                index = index + 1
        elif data:
            _id = 1
            for items in data:
                self.doc.add_paragraph(style="Heading 2").add_run(str(_id) + "." + items['name'],
                                                                  style="Default Paragraph Font")
                _id = _id + 1
                self.doc.add_paragraph("简要描述：", style="Body Text").paragraph_format.space_before = Pt(14)
                self.doc.add_paragraph(items['description'], style="List Paragraph")
                self.doc.add_paragraph("请求URL：", style="Body Text")
                if items['httpType'] == 'HTTP':
                    self.doc.add_paragraph("http://xxxx/" + items['apiAddress'], style="List Paragraph")
                else:
                    self.doc.add_paragraph("https://xxxx/" + items['apiAddress'], style="List Paragraph")
                self.doc.add_paragraph("请求方式：", style="Body Text")
                self.doc.add_paragraph(items['requestType'], style="List Paragraph")
                self.doc.add_paragraph("参数：", style="Body Text")
                if items['requestParameterType'] == 'form-data':
                    table = self.doc.add_table(rows=len(items['requestParameter']) + 1,
                                               cols=4, style="Medium Shading 1 Accent 1")
                    hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
                    hdr_cells[0].text = '参数名'  # 第一行的第一列,给这行里面添加文字
                    hdr_cells[1].text = '参数类型'
                    hdr_cells[2].text = '必填?'
                    hdr_cells[3].text = '输入限制'
                    j = 0
                    for row in table.rows[1:]:
                        row.cells[0].text = items['requestParameter'][j]['name']
                        row.cells[1].text = items['requestParameter'][j]['_type']
                        if items['requestParameter'][j]['required']:
                            row.cells[2].text = "是"
                        else:
                            row.cells[2].text = "否"
                        try:
                            row.cells[3].text = items['requestParameter'][j]['restrict']
                        except TypeError:
                            pass
                        j = j + 1
                else:
                    try:
                        if len(items['requestParameterRaw']):
                            data = json.loads(items['requestParameterRaw'][0]["data"])
                            self.doc.add_paragraph(style="Normal").add_run('{')
                            write_json(self.doc, data, 0.3)
                            self.doc.add_paragraph(style="Normal").add_run('}')
                    except:
                        logging.exception("Error")
                self.doc.add_paragraph("返回参数：", style="Body Text")
                table = self.doc.add_table(rows=len(items['response']) + 1,
                                           cols=4, style="Medium Shading 1 Accent 1")
                hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
                hdr_cells[0].text = '参数名'  # 第一行的第一列,给这行里面添加文字
                hdr_cells[1].text = '参数类型'
                hdr_cells[2].text = '必含?'
                hdr_cells[3].text = '说明'
                j = 0
                for row in table.rows[1:]:
                    row.cells[0].text = items['response'][j]['name']
                    row.cells[1].text = items['response'][j]['_type']
                    if items['response'][j]['required']:
                        row.cells[2].text = "是"
                    else:
                        row.cells[2].text = "否"
                    try:
                        row.cells[3].text = items['response'][j]['description']
                    except TypeError:
                        pass
                    j = j + 1
                self.doc.add_paragraph()
                self.doc.add_paragraph("返回示例：", style="Body Text")
                try:
                    if len(items['data']):
                        # data = eval(items['data'].replace("true", "True").replace("false", "False"))
                        data = json.loads(items["data"])
                        self.doc.add_paragraph(style="Normal").add_run('{')
                        write_json(self.doc, data, 0.3)
                        self.doc.add_paragraph(style="Normal").add_run('}')
                except:
                    logging.exception("Error")
        path = "./api_test/ApiDoc/%s.docx" % api_name
        self.doc.save(path)
        return path


def write_json(doc, data, num):
    try:
        for n in data:
            p = doc.add_paragraph(style="Normal")
            run = p.add_run('"%s":' % n)
            run.font.color.rgb = RGBColor(186, 85, 211)
            p.paragraph_format.first_line_indent = Inches(num)
            if isinstance(data[n], dict):
                write_json(doc, data[n], num+0.3)
            else:
                if data[n] is None:
                    run = p.add_run(' null,')
                    run.font.color.rgb = RGBColor(186, 85, 211)
                elif isinstance(data[n], bool):
                    run = p.add_run(' %s,' % data[n])
                    run.font.color.rgb = RGBColor(255, 0, 0)
                elif isinstance(data[n], int):
                    run = p.add_run(' %s,' % data[n])
                    run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run = p.add_run(' "%s",' % data[n])
                    run.font.color.rgb = RGBColor(59, 149, 38)
    except:
        logging.exception("error")
